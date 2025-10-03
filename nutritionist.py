import streamlit as st
import chromadb
from chromadb.config import Settings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_ollama import OllamaLLM
from langchain.chains.retrieval_qa.base import RetrievalQA
from langchain_core.prompts import PromptTemplate
from langchain_core.documents import Document
import PyPDF2
import io
import json
from datetime import datetime
import docx
import pandas as pd

# Page configuration with dark theme
st.set_page_config(
    page_title="Analyze Your Diet",
    page_icon="🍏",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for DeepSeek-like styling
st.markdown("""
<style>
    /* Dark theme */
    .stApp {
        background-color: #0a0a0a;
        color: #e0e0e0;
    }
    
    /* Hide default Streamlit elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Custom header */
    .custom-header {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
        padding: 1.5rem 2rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        border: 1px solid #2a2a3e;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.3);
    }
    
    .custom-header h1 {
        color: #ffffff;
        font-size: 2rem;
        font-weight: 600;
        margin: 0;
        display: flex;
        align-items: center;
        gap: 0.75rem;
    }
    
    .custom-header p {
        color: #a0a0a0;
        margin: 0.5rem 0 0 0;
        font-size: 0.95rem;
    }
    
    /* Chat-like container */
    .chat-container {
        max-width: 900px;
        margin: 0 auto;
        padding: 0 1rem;
    }
    
    /* Message bubbles */
    .user-message {
        background: linear-gradient(135deg, #1e3a8a 0%, #1e40af 100%);
        padding: 1.25rem 1.5rem;
        border-radius: 18px;
        margin: 1rem 0;
        border: 1px solid #2563eb;
        box-shadow: 0 2px 8px rgba(37, 99, 235, 0.2);
    }
    
    .ai-message {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
        padding: 1.5rem;
        border-radius: 18px;
        margin: 1rem 0;
        border: 1px solid #2a2a3e;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.3);
    }
    
    .ai-message h3 {
        color: #60a5fa;
        font-size: 1.1rem;
        margin-bottom: 1rem;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
    
    /* Buttons */
    .stButton > button {
        background: linear-gradient(135deg, #2563eb 0%, #1d4ed8 100%);
        color: white;
        border: none;
        border-radius: 12px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        font-size: 1rem;
        width: 100%;
        transition: all 0.3s ease;
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.3);
    }
    
    .stButton > button:hover {
        background: linear-gradient(135deg, #1d4ed8 0%, #1e40af 100%);
        box-shadow: 0 6px 16px rgba(37, 99, 235, 0.4);
        transform: translateY(-2px);
    }
    
    /* File uploader */
    .uploadedFile {
        background: #1a1a2e;
        border: 2px dashed #2a2a3e;
        border-radius: 12px;
        padding: 2rem;
    }
    
    /* Metrics */
    .metric-card {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
        padding: 1.5rem;
        border-radius: 16px;
        border: 1px solid #2a2a3e;
        text-align: center;
        margin: 0.5rem;
    }
    
    .metric-value {
        font-size: 2.5rem;
        font-weight: 700;
        background: linear-gradient(135deg, #60a5fa 0%, #3b82f6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin: 0.5rem 0;
    }
    
    .metric-label {
        color: #a0a0a0;
        font-size: 0.9rem;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }
    
    /* Tags */
    .food-tag {
        display: inline-block;
        background: #1e3a8a;
        color: #93c5fd;
        padding: 0.4rem 1rem;
        border-radius: 20px;
        margin: 0.25rem;
        font-size: 0.85rem;
        border: 1px solid #2563eb;
    }
    
    /* Status badge */
    .status-badge {
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        background: #065f46;
        color: #6ee7b7;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 600;
        border: 1px solid #059669;
    }
    
    /* Loading spinner */
    .stSpinner > div {
        border-color: #3b82f6 !important;
    }
    
    /* Expander */
    .streamlit-expanderHeader {
        background: #1a1a2e;
        border-radius: 12px;
        border: 1px solid #2a2a3e;
        color: #e0e0e0 !important;
    }
    
    /* Info boxes */
    .stAlert {
        background: #1a1a2e;
        border: 1px solid #2a2a3e;
        border-radius: 12px;
        color: #e0e0e0;
    }
    
    /* Progress bar */
    .stProgress > div > div {
        background: linear-gradient(90deg, #2563eb 0%, #3b82f6 100%);
    }
    
    /* Input fields */
    .stTextInput > div > div > input {
        background: #1a1a2e;
        border: 1px solid #2a2a3e;
        border-radius: 12px;
        color: #e0e0e0;
        padding: 0.75rem 1rem;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #3b82f6;
        box-shadow: 0 0 0 2px rgba(59, 130, 246, 0.2);
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'chroma_db' not in st.session_state:
    st.session_state.chroma_db = None
if 'qa_chain' not in st.session_state:
    st.session_state.qa_chain = None
if 'analysis_history' not in st.session_state:
    st.session_state.analysis_history = []
if 'nutrition_knowledge_loaded' not in st.session_state:
    st.session_state.nutrition_knowledge_loaded = False
if 'ollama_url' not in st.session_state:
    st.session_state.ollama_url = "http://localhost:11434"

# Nutritional knowledge base
NUTRITION_KNOWLEDGE = """
# Nutritional Guidelines and Information

## Macronutrients
- Proteins: Essential for muscle repair, enzyme production, and immune function. Adults need 0.8g per kg of body weight daily.
- Carbohydrates: Primary energy source. Complex carbs (whole grains, vegetables) are preferred over simple sugars.
- Fats: Essential for hormone production and nutrient absorption. Focus on unsaturated fats from nuts, avocados, and fish.

## Micronutrients
- Vitamins: A, B-complex, C, D, E, K each play unique roles in health.
- Minerals: Calcium, iron, magnesium, zinc, potassium are crucial for various bodily functions.

## Food Groups
1. Vegetables: 5+ servings daily. Rich in fiber, vitamins, minerals, and antioxidants.
2. Fruits: 2-3 servings daily. Natural sugars with fiber, vitamins, and phytonutrients.
3. Proteins: Lean meats, fish, eggs, legumes, tofu. Essential for tissue repair.
4. Whole Grains: Brown rice, quinoa, oats. Provide sustained energy and fiber.
5. Dairy/Alternatives: Calcium and vitamin D for bone health.
6. Healthy Fats: Nuts, seeds, avocados, olive oil. Support brain and heart health.

## Dietary Patterns
- Mediterranean Diet: Rich in vegetables, fruits, whole grains, fish, olive oil.
- DASH Diet: Focuses on reducing sodium, increasing potassium-rich foods.
- Plant-Based: Emphasizes whole plant foods, may include or exclude animal products.

## Common Deficiencies
- Vitamin D: Common in northern climates, important for immune function.
- Iron: Particularly in menstruating women, vegetarians.
- B12: Risk for vegans, older adults.
- Omega-3: Important for heart and brain health.

## Processed Foods
- Ultra-processed foods often contain high sodium, sugar, unhealthy fats, and additives.
- Limit consumption of packaged snacks, sugary drinks, fast food.
- Focus on whole, minimally processed foods.

## Hydration
- Water: 8-10 glasses daily, more with exercise.
- Limit sugary drinks, excessive caffeine.

## Meal Timing
- Regular meal patterns support metabolism.
- Breakfast importance varies by individual.
- Evening eating should be balanced, not excessive.

## Special Considerations
- Athletes: Higher protein and calorie needs.
- Pregnancy: Increased folate, iron, calcium needs.
- Aging: May need more protein, calcium, vitamin D.
- Diabetes: Carbohydrate counting and timing important.
"""

def initialize_rag_system(ollama_url="http://localhost:11434"):
    """Initialize the RAG system with ChromaDB and LangChain"""
    try:
        embeddings = OllamaEmbeddings(
            base_url=ollama_url,
            model="llama3.2:3b"
        )
        
        chroma_client = chromadb.Client(Settings(
            anonymized_telemetry=False,
            allow_reset=True
        ))
        
        try:
            chroma_client.delete_collection("nutrition_knowledge")
        except:
            pass
        
        collection = chroma_client.create_collection(
            name="nutrition_knowledge",
            metadata={"description": "Nutritional information and guidelines"}
        )
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len
        )
        
        documents = [Document(page_content=NUTRITION_KNOWLEDGE)]
        splits = text_splitter.split_documents(documents)
        
        vectorstore = Chroma.from_documents(
            documents=splits,
            embedding=embeddings,
            collection_name="nutrition_knowledge",
            client=chroma_client
        )
        
        llm = OllamaLLM(
            base_url=ollama_url,
            model="llama3.2:3b",
            temperature=0.7
        )
        
        prompt_template = """You are an expert nutritionist analyzing diet data. Use the following nutritional information to provide accurate, evidence-based advice.

Context: {context}

Question: {question}

Provide a detailed, professional analysis that includes:
1. Assessment of the current diet based on nutritional guidelines
2. Specific strengths and areas for improvement
3. Evidence-based recommendations
4. Potential health implications

Answer:"""

        PROMPT = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )
        
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
            chain_type_kwargs={"prompt": PROMPT},
            return_source_documents=True
        )
        
        return vectorstore, qa_chain, True
        
    except Exception as e:
        st.error(f"Error initializing system: {str(e)}")
        return None, None, False

def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.read()))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None

def extract_text_from_txt(txt_file):
    """Extract text from uploaded TXT file"""
    try:
        text = txt_file.read().decode('utf-8')
        return text
    except Exception as e:
        st.error(f"Error reading TXT file: {str(e)}")
        return None

def extract_text_from_docx(docx_file):
    """Extract text from uploaded DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(docx_file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX file: {str(e)}")
        return None

def extract_text_from_csv(csv_file):
    """Extract text from uploaded CSV file"""
    try:
        df = pd.read_csv(io.BytesIO(csv_file.read()))
        # Convert all columns to string and combine
        text = " ".join(df.astype(str).values.flatten())
        return text
    except Exception as e:
        st.error(f"Error reading CSV file: {str(e)}")
        return None

def extract_text_from_file(uploaded_file):
    """Extract text from various file formats"""
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(uploaded_file)
    elif file_extension == 'txt':
        return extract_text_from_txt(uploaded_file)
    elif file_extension == 'docx':
        return extract_text_from_docx(uploaded_file)
    elif file_extension == 'csv':
        return extract_text_from_csv(uploaded_file)
    else:
        st.error(f"Unsupported file format: {file_extension}")
        return None

def categorize_foods(food_list):
    """Categorize foods into groups"""
    categories = {
        'vegetables': ['lettuce', 'carrot', 'broccoli', 'spinach', 'kale', 'tomato', 
                      'cucumber', 'pepper', 'celery', 'cabbage', 'salad', 'greens', 'beans',
                      'cauliflower', 'zucchini', 'eggplant', 'mushroom', 'onion', 'garlic',
                      'asparagus', 'brussels', 'sprouts', 'potato', 'sweet potato', 'yam'],
        'fruits': ['apple', 'banana', 'orange', 'berry', 'berries', 'grape', 'melon', 
                  'peach', 'pear', 'mango', 'strawberry', 'blueberry', 'raspberry',
                  'pineapple', 'kiwi', 'cherry', 'plum', 'apricot', 'fig', 'date',
                  'watermelon', 'cantaloupe', 'honeydew', 'pomegranate', 'coconut'],
        'proteins': ['chicken', 'beef', 'pork', 'fish', 'salmon', 'tuna', 'egg', 
                    'tofu', 'turkey', 'shrimp', 'meat', 'lentil', 'bean', 'tempeh',
                    'seitan', 'lamb', 'duck', 'bison', 'venison', 'cod', 'tilapia',
                    'halibut', 'sardine', 'anchovy', 'crab', 'lobster', 'scallop',
                    'clam', 'mussel', 'chickpea', 'black bean', 'kidney bean'],
        'grains': ['rice', 'bread', 'pasta', 'cereal', 'oats', 'quinoa', 'wheat', 
                  'bagel', 'tortilla', 'noodle', 'barley', 'millet', 'buckwheat',
                  'couscous', 'bulgur', 'farro', 'spelt', 'rye', 'corn', 'popcorn'],
        'dairy': ['milk', 'cheese', 'yogurt', 'butter', 'cream', 'kefir', 'cottage',
                 'sour cream', 'whipped cream', 'ice cream', 'ghee', 'whey'],
        'processed': ['pizza', 'burger', 'fries', 'chips', 'candy', 'soda', 
                     'cookie', 'cake', 'donut', 'hot dog', 'nugget', 'sausage',
                     'bacon', 'ham', 'lunchmeat', 'frozen dinner', 'instant noodle',
                     'white bread', 'sugar', 'syrup', 'ketchup', 'mayonnaise'],
        'healthy_fats': ['avocado', 'nuts', 'seeds', 'olive', 'almond', 'walnut',
                        'pecan', 'cashew', 'pistachio', 'macadamia', 'hazelnut',
                        'sunflower seed', 'pumpkin seed', 'chia seed', 'flaxseed',
                        'coconut oil', 'avocado oil', 'nut butter', 'tahini']
    }
    
    category_counts = {cat: 0 for cat in categories.keys()}
    categorized_items = {cat: [] for cat in categories.keys()}
    
    for food in food_list:
        food_lower = food.lower()
        categorized = False
        for category, keywords in categories.items():
            if any(keyword in food_lower for keyword in keywords):
                category_counts[category] += 1
                categorized_items[category].append(food)
                categorized = True
                break
        
        # If not categorized, add to "other" category
        if not categorized:
            if 'other' not in category_counts:
                category_counts['other'] = 0
                categorized_items['other'] = []
            category_counts['other'] += 1
            categorized_items['other'].append(food)
    
    return category_counts, categorized_items

def calculate_health_score(category_counts):
    """Calculate health score based on food categories"""
    healthy = category_counts['vegetables'] + category_counts['fruits'] + category_counts['healthy_fats']
    unhealthy = category_counts['processed']
    protein = category_counts['proteins']
    
    score = min(100, max(0, (healthy * 10) - (unhealthy * 5) + (protein * 5) + 30))
    return round(score)

def analyze_diet_with_rag(qa_chain, food_data, category_counts, health_score):
    """Use RAG to analyze diet with retrieved nutritional knowledge"""
    
    query = f"""
    Analyze this monthly diet data:
    
    Total food items: {len(food_data)}
    
    Category breakdown:
    - Vegetables: {category_counts['vegetables']} items
    - Fruits: {category_counts['fruits']} items
    - Proteins: {category_counts['proteins']} items
    - Grains: {category_counts['grains']} items
    - Dairy: {category_counts['dairy']} items
    - Processed foods: {category_counts['processed']} items
    - Healthy fats: {category_counts['healthy_fats']} items
    - Other: {category_counts.get('other', 0)} items
    
    Current health score: {health_score}/100
    
    Sample foods: {', '.join(food_data[:15])}
    
    Provide a comprehensive nutritional analysis with specific recommendations.
    """
    
    try:
        result = qa_chain({"query": query})
        return result['result'], result.get('source_documents', [])
    except Exception as e:
        return f"Error during analysis: {str(e)}", []

def add_documents_to_vectorstore(vectorstore, documents, embeddings):
    """Add user's diet documents to the vector store for context"""
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=300,
            chunk_overlap=30
        )
        splits = text_splitter.split_documents(documents)
        vectorstore.add_documents(splits)
        return True
    except Exception as e:
        st.error(f"Error adding documents: {str(e)}")
        return False

# Main App
def main():
    st.markdown("""
        <h1>Diet Analyzer</h1>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    st.markdown('<div class="ai-message">', unsafe_allow_html=True)
    st.markdown("### System Configuration")
    
    col1, col2 = st.columns([3, 1])
    with col1:
        ollama_url = st.text_input(
            "Ollama API URL",
            value=st.session_state.ollama_url,
            help="Default: http://localhost:11434"
        )
        st.session_state.ollama_url = ollama_url
    
    with col2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("Initialize", use_container_width=True):
            with st.spinner("Initializing agent..."):
                vectorstore, qa_chain, success = initialize_rag_system(ollama_url)
                if success:
                    st.session_state.chroma_db = vectorstore
                    st.session_state.qa_chain = qa_chain
                    st.session_state.nutrition_knowledge_loaded = True
                    st.success("✅ System ready!")
                    st.rerun()
    
    if st.session_state.nutrition_knowledge_loaded:
        st.markdown("""
            <div style="margin-top: 1rem;">
                <span class="status-badge"> ● System Active</span>
            </div>
        """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Main interaction area
    st.markdown("<br>", unsafe_allow_html=True)
    
    # File upload in a message-like container
    st.markdown('<div class="user-message">', unsafe_allow_html=True)
    st.markdown("### 📄 Upload Your Files")
    uploaded_file = st.file_uploader(
        "Upload your food diary or diet log",
        type=['pdf', 'txt', 'docx', 'csv'],
        help="Supported formats: PDF, TXT, DOCX, CSV",
        label_visibility="collapsed"
    )
    st.markdown('</div>', unsafe_allow_html=True)
    
    if uploaded_file:
        with st.spinner("Processing..."):
            file_text = extract_text_from_file(uploaded_file)
            
            if file_text:
                words = file_text.lower().split()
                # Enhanced stop words list
                stop_words = {'the', 'and', 'with', 'for', 'week', 'day', 'month', 'year',
                             'breakfast', 'lunch', 'dinner', 'snack', 'meal', 'food',
                             'this', 'that', 'these', 'those', 'have', 'had', 'has',
                             'eat', 'ate', 'eating', 'drink', 'drank', 'drinking',
                             'morning', 'afternoon', 'evening', 'night', 'today',
                             'yesterday', 'tomorrow', 'monday', 'tuesday', 'wednesday',
                             'thursday', 'friday', 'saturday', 'sunday', 'january',
                             'february', 'march', 'april', 'may', 'june', 'july',
                             'august', 'september', 'october', 'november', 'december'}
                
                food_keywords = set(words) - stop_words
                # Filter out very short words (likely not food items)
                foods = [word for word in list(food_keywords) if len(word) > 2][:50]
                
                st.session_state.foods = foods
                st.session_state.file_text = file_text
                
                # Show extracted data in AI message style
                st.markdown('<div class="ai-message">', unsafe_allow_html=True)
                st.markdown("### ✅ Document Processed")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.markdown(f"""
                        <div class="metric-card">
                            <div class="metric-label">Food Items</div>
                            <div class="metric-value">{len(foods)}</div>
                        </div>
                    """, unsafe_allow_html=True)
                
                with col2:
                    st.markdown(f"""
                        <div class="metric-card">
                            <div class="metric-label">Characters</div>
                            <div class="metric-value">{len(file_text)}</div>
                        </div>
                    """, unsafe_allow_html=True)
                
                with col3:
                    category_counts, _ = categorize_foods(foods)
                    health_score = calculate_health_score(category_counts)
                    st.markdown(f"""
                        <div class="metric-card">
                            <div class="metric-label">Preview Score</div>
                            <div class="metric-value">{health_score}</div>
                        </div>
                    """, unsafe_allow_html=True)
                
                st.markdown("<br>", unsafe_allow_html=True)
                st.markdown(f"**File type:** {uploaded_file.name.split('.')[-1].upper()}")
                st.markdown("**Sample extracted foods:**")
                st.markdown(" ".join([f'<span class="food-tag">{food}</span>' for food in foods[:15]]), unsafe_allow_html=True)
                
                st.markdown('</div>', unsafe_allow_html=True)
    
    # Analysis button
    if 'foods' in st.session_state and st.session_state.nutrition_knowledge_loaded:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("Analyze My Diet", use_container_width=True):
            with st.spinner("Analyzing..."):
                foods = st.session_state.foods
                category_counts, categorized = categorize_foods(foods)
                health_score = calculate_health_score(category_counts)
                
                # Add to vector store
                diet_doc = Document(
                    page_content=f"User's diet contains: {', '.join(foods)}. {st.session_state.file_text[:1000]}"
                )
                add_documents_to_vectorstore(
                    st.session_state.chroma_db,
                    [diet_doc],
                    OllamaEmbeddings(base_url=st.session_state.ollama_url, model="llama3.2:3b")
                )
                
                # Perform analysis
                analysis, sources = analyze_diet_with_rag(
                    st.session_state.qa_chain,
                    foods,
                    category_counts,
                    health_score
                )
                
                st.session_state.analysis_history.append({
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'health_score': health_score,
                    'analysis': analysis,
                    'categories': category_counts
                })
                
                # Display results
                st.markdown("<br>", unsafe_allow_html=True)
                
                # Health score in prominent display
                st.markdown('<div class="ai-message">', unsafe_allow_html=True)
                st.markdown("### 📊 Your Health Score")
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.markdown(f"""
                        <div class="metric-card">
                            <div class="metric-label">Overall Score</div>
                            <div class="metric-value">{health_score}</div>
                        </div>
                    """, unsafe_allow_html=True)
                
                with col2:
                    status = "Excellent" if health_score > 70 else "Good" if health_score > 40 else "Needs Work"
                    st.markdown(f"""
                        <div class="metric-card">
                            <div class="metric-label">Status</div>
                            <div class="metric-value" style="font-size: 1.5rem;">{status}</div>
                        </div>
                    """, unsafe_allow_html=True)
                
                with col3:
                    healthy_items = category_counts['vegetables'] + category_counts['fruits']
                    st.markdown(f"""
                        <div class="metric-card">
                            <div class="metric-label">Healthy Foods</div>
                            <div class="metric-value">{healthy_items}</div>
                        </div>
                    """, unsafe_allow_html=True)
                
                with col4:
                    st.markdown(f"""
                        <div class="metric-card">
                            <div class="metric-label">Processed</div>
                            <div class="metric-value">{category_counts['processed']}</div>
                        </div>
                    """, unsafe_allow_html=True)
                
                st.markdown('</div>', unsafe_allow_html=True)
                
                # AI Analysis
                st.markdown("<br>", unsafe_allow_html=True)
                st.markdown('<div class="ai-message">', unsafe_allow_html=True)
                st.markdown("### Nutritional Analysis")
                st.markdown(analysis)
                st.markdown('</div>', unsafe_allow_html=True)
                
                # Category breakdown
                st.markdown("<br>", unsafe_allow_html=True)
                st.markdown('<div class="ai-message">', unsafe_allow_html=True)
                st.markdown("### 📋 Detailed Breakdown")
                
                # Include 'other' category in display
                display_categories = {k: v for k, v in category_counts.items() if v > 0}
                categories_list = list(display_categories.items())
                
                # Create dynamic columns based on number of categories
                num_cols = min(4, len(categories_list))
                cols = st.columns(num_cols)
                
                for idx, (category, count) in enumerate(categories_list):
                    with cols[idx % num_cols]:
                        st.markdown(f"""
                            <div style="text-align: center; padding: 1rem;">
                                <div style="color: #a0a0a0; font-size: 0.85rem; text-transform: uppercase;">{category.replace('_', ' ')}</div>
                                <div style="color: #60a5fa; font-size: 1.8rem; font-weight: 700; margin-top: 0.5rem;">{count}</div>
                            </div>
                        """, unsafe_allow_html=True)
                
                st.markdown('</div>', unsafe_allow_html=True)
                
                # Detailed foods by category
                with st.expander("🔍 View All Foods by Category"):
                    for category, items in categorized.items():
                        if items:
                            st.markdown(f"**{category.replace('_', ' ').title()}** ({len(items)} items)")
                            st.markdown(" ".join([f'<span class="food-tag">{item}</span>' for item in items[:20]]), unsafe_allow_html=True)
                            if len(items) > 20:
                                st.markdown(f"<p style='color: #a0a0a0; font-size: 0.85rem;'>... and {len(items) - 20} more</p>", unsafe_allow_html=True)
                            st.markdown("<br>", unsafe_allow_html=True)
    
    elif 'foods' in st.session_state and not st.session_state.nutrition_knowledge_loaded:
        st.warning("⚠️ Please initialize the system first using the settings above.")
    
    # Analysis history
    if st.session_state.analysis_history:
        st.markdown("<br><br>", unsafe_allow_html=True)
        with st.expander("📜 Previous Analyses"):
            for idx, record in enumerate(reversed(st.session_state.analysis_history[-5:])):
                st.markdown(f"**{record['timestamp']}** • Score: {record['health_score']}/100")
                st.text(record['analysis'][:300] + "...")
                st.markdown("---")
    
    st.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()